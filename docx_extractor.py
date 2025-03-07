import docx
import os
import logging
from PIL import Image, ImageDraw, ImageFont
import re
import math

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    """Extract text from a Word document using python-docx."""
    try:
        logger.info(f"Extracting text from DOCX: {file_path}")
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Skip empty cells
                        full_text.append(cell.text)
        
        # Count the number of pages (approximate since python-docx doesn't provide page count)
        # A rough estimate based on content size
        page_count = max(1, len('\n'.join(full_text)) // 3000)
        
        logger.info(f"Successfully extracted text from DOCX: {file_path}")
        return '\n'.join(full_text), page_count
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise

def create_docx_preview_info(file_path, estimated_page_count=None):
    """Create preview information without generating any preview images.
    
    Args:
        file_path: Path to the DOCX file
        estimated_page_count: Optional estimate of page count
        
    Returns:
        dict: Dictionary containing preview information
    """
    try:
        logger.info(f"Creating preview info for DOCX: {file_path}")
        
        # Get the base filename without extension
        base_filename = os.path.splitext(os.path.basename(file_path))[0]
        
        # For DOCX, estimate the number of sections based on page count
        # For 1-5 pages, use 1 section per page
        # For 6+ pages, use 1 section per 2 pages (approximately)
        if estimated_page_count is None:
            try:
                # Estimate page count from content
                doc = docx.Document(file_path)
                all_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                estimated_page_count = max(1, len(all_text) // 3000)
            except:
                estimated_page_count = 5  # Default if we can't estimate
        
        # Calculate number of sections
        if estimated_page_count <= 5:
            section_count = estimated_page_count
        else:
            section_count = max(5, estimated_page_count // 2)
        
        # Create preview info
        preview_info = {
            'base_filename': base_filename,
            'document_page_count': section_count,  # For DOCX, we use sections not actual pages
            'preview_format': "{base}_section_{page}.png",
            'file_type': 'docx'
        }
        
        return preview_info
    except Exception as e:
        logger.error(f"Error creating DOCX preview info for {file_path}: {str(e)}")
        return None

def generate_section_preview(document, section_number, preview_dir):
    """Generate a preview image for a specific section of a DOCX document.
    
    Args:
        document: Document object
        section_number: Section number to generate (1-based index)
        preview_dir: Directory to save preview image
        
    Returns:
        str: Filename of the generated preview image, or None if failed
    """
    try:
        # Get preview info
        preview_info = document.get_preview_info()
        if not preview_info:
            logger.error(f"No preview info found for document ID {document.id}")
            return None
            
        # Check if section number is valid
        if section_number < 1 or section_number > preview_info.get('document_page_count', 0):
            logger.error(f"Invalid section number {section_number} for document ID {document.id}")
            return None
            
        # Construct the expected preview filename
        preview_filename = preview_info['preview_format'].format(
            base=preview_info['base_filename'],
            page=section_number
        )
        preview_path = os.path.join(preview_dir, preview_filename)
        
        # Check if preview already exists
        if os.path.exists(preview_path):
            logger.info(f"Preview for section {section_number} already exists at: {preview_path}")
            return preview_filename
            
        # Open the DOCX document
        file_path = os.path.join(os.path.dirname(preview_dir), document.filename)
        if not os.path.exists(file_path):
            logger.error(f"Document file not found: {file_path}")
            return None
            
        doc = docx.Document(file_path)
        
        # Extract title and properties
        title = preview_info['base_filename']
        try:
            # Try to get document title from properties
            core_props = doc.core_properties
            if core_props.title:
                title = core_props.title
        except:
            pass
            
        # Clean up the title
        title = re.sub(r'[_-]', ' ', title)
        
        # Extract all paragraphs with content
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract all tables
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(" | ".join(row_text))
            tables.append("\n".join(table_text))
        
        # Calculate how many "sections" to generate
        content_items = paragraphs + tables
        total_items = len(content_items)
        
        if total_items == 0:
            # No content to preview, create at least one page
            total_items = 1
            content_items = ["No content available in document"]
            
        # Divide content items into sections
        sections_count = preview_info.get('document_page_count', 1)
        items_per_section = math.ceil(total_items / sections_count)
        
        # Calculate which content items to include in this section
        start_idx = (section_number - 1) * items_per_section
        end_idx = min(total_items, start_idx + items_per_section)
        
        # Make sure we have a valid range
        if start_idx >= total_items:
            start_idx = 0
            end_idx = min(total_items, items_per_section)
            
        section_items = content_items[start_idx:end_idx]
        
        # Create a blank image (letter size 8.5x11 inches at 100 DPI)
        width, height = 850, 1100  # 8.5x11 inches at 100 DPI
        image = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(image)
        
        # Try to get a decent font - fallback to default if necessary
        try:
            # Try to find a system font
            font_large = ImageFont.truetype("Arial", 36)
            font_medium = ImageFont.truetype("Arial", 24)
            font_small = ImageFont.truetype("Arial", 16)
        except IOError:
            try:
                # Try DejaVuSans as an alternative on Linux/Unix systems
                font_large = ImageFont.truetype("DejaVuSans", 36)
                font_medium = ImageFont.truetype("DejaVuSans", 24)
                font_small = ImageFont.truetype("DejaVuSans", 16)
            except IOError:
                # Fallback to default
                font_large = ImageFont.load_default()
                font_medium = font_large
                font_small = font_large
        
        # Draw a border
        draw.rectangle([20, 20, width-20, height-20], outline='black', width=2)
        
        # Draw the top header box
        draw.rectangle([20, 20, width-20, 140], fill='lightblue', outline='black', width=2)
        
        # Draw the document title
        draw.text((width//2, 80), title, fill='black', font=font_large, anchor='mm')
        
        # Draw section label
        section_label = f"Section {section_number} of {sections_count}"
        draw.text((width//2, 120), section_label, fill='black', font=font_medium, anchor='mm')
        
        # Draw document content for this section
        y_pos = 180
        draw.text((40, y_pos), f"Document Content - {end_idx-start_idx} items (Items {start_idx+1}-{end_idx})", fill='black', font=font_medium)
        y_pos += 50
        
        # Draw paragraphs/content for this section
        for i, text in enumerate(section_items):
            if len(text) > 80:  # Truncate long paragraphs
                text = text[:77] + "..."
            
            draw.text((40, y_pos), text, fill='black', font=font_small)
            y_pos += 30
            
            if y_pos > height - 100:  # Stop if we're running out of space
                draw.text((40, y_pos), "... more content ...", fill='black', font=font_small)
                break
        
        # Draw footer
        draw.rectangle([20, height-70, width-20, height-20], fill='lightblue', outline='black', width=2)
        draw.text((width//2, height-45), "Document Analyzer Preview", fill='black', font=font_small, anchor='mm')
        
        # Save the image
        image.save(preview_path)
        
        logger.info(f"Preview generated for section {section_number} at: {preview_path}")
        return preview_filename
    except Exception as e:
        logger.error(f"Error generating section preview: {str(e)}")
        return None